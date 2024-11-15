from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from data import *

# Crear el documento
doc = Document()

section = doc.sections[0]
section.left_margin = Cm(Margenes.izquierdo)    # Margen izquierdo
section.right_margin = Cm(Margenes.derecho)   # Margen derecho
section.top_margin = Cm(Margenes.superior)     # Margen superior
section.bottom_margin = Cm(Margenes.inferior)  # Margen inferior

# Agregar un encabezado
section = doc.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.text = Nombre_Ano
header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Estilo del encabezado
header_run = header_paragraph.runs[0]
header_run.font.size = Pt(10)
header_run.font.name = 'Arial'
header_run.font.bold = True

# Configurar el estilo global del documento
estilo = doc.styles['Normal']
fuente = estilo.font
fuente.name = 'Arial'  # Tipo de letra
fuente.size = Pt(10)   # Tamaño de letra


#Agregar espacio
Espacio = doc.add_paragraph()
Espacio_red = Espacio.add_run("")

# Agregar fecha y hora de la carta
fecha_hora_parrafo = doc.add_paragraph()
fecha_hora = fecha_hora_parrafo.add_run(f'{fecha_departamento.departamento}, {fecha_departamento.fecha} \n\r')

# titulo de la carta
titulo_parrafo = doc.add_paragraph()
titulo = titulo_parrafo.add_run(f'CARTA N° {emisor.n_carta}-{emisor.ano}-{emisor.iniciales.upper()}')
titulo.font.bold = True
titulo.font.underline = True

# Emisor
parrafo = doc.add_paragraph()
genero_receptor = parrafo.add_run(f'{receptor.genero}:\n')

persona_receptor = parrafo.add_run(f'{receptor.director.upper()}\n')
persona_receptor.font.bold = True

direccion_receprtor = parrafo.add_run(f'{receptor.direccion}\n')

lugar_receptor = parrafo.add_run('Organismo de Evaluación y Fiscalización Ambiental-OEFA Av.\nFaustino Sánchez Carrión N° 603 - Jesús María')

#Formalismo
parrafo = doc.add_paragraph()
texto_atencion = parrafo.add_run("Atención")
texto_atencion.underline = True     # Subrayado para "Atención"

texto_subdireccion = parrafo.add_run(": Subdirección Técnica Científica")

parrafo = doc.add_paragraph()
texto_presente = parrafo.add_run("Presente")
texto_presente.underline = True     # Subrayado para "Atención"

texto_subdireccion = parrafo.add_run(": -")

parrafo = doc.add_paragraph()
texto_asunto = parrafo.add_run("Asunto")
texto_presente.underline = True     # Subrayado para "Atención"

if "N°" in emisor.Adenda:
    texto_contrato = f"{emisor.entregable} entregable en el marco de la Adenda {emisor.Adenda} del Contrato {emisor.contrato}"

else:

    texto_contrato = f"{emisor.entregable} entregable en el marco del Contrato {emisor.contrato}"


texto_subdireccion = parrafo.add_run(f": Presentación del {texto_contrato}")
parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 

parrafo = doc.add_paragraph()
cuerpo_text = parrafo.add_run("De mi consideración:")

parrafo = doc.add_paragraph()
cuerpo2_text = parrafo.add_run(f"Me dirijo a usted, para expresarle mi cordial saludo, y a la vez adjuntar el {texto_contrato},")
cuerpo3_text = parrafo.add_run(f"de acuerdo con el Detalle de Actividades solicitado por la Subdirección Técnica Científica de la Dirección de Evaluación Ambiental.")
parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 

parrafo = doc.add_paragraph()
cuerpo4_text = parrafo.add_run("Me despido cordialmente agradeciéndole la atención a la presente.")

parrafo = doc.add_paragraph()
cuerpo4_text = parrafo.add_run("Atentamente,")

parrafo = doc.add_paragraph()
cuerpo5_text = parrafo.add_run("\n\r\n\r\n\r\n\r\n\r\n\r")
cuerpo6_text = parrafo.add_run("………………………………………………")
cuerpo7_text = parrafo.add_run(f"\n{emisor.nombre}")
cuerpo8_text = parrafo.add_run(f"\nDNI:{emisor.dni}")
cuerpo9_text = parrafo.add_run("\n\rSe Adjunta:")
cuerpo10_text = parrafo.add_run("\n\t-\tInforme de Actividades y Anexos")

# Guardar el documento
doc.save("dodoc.docx")