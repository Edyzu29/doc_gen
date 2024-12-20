from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from data import *
from detalle_data import *

# Crear el documento
doc = Document()

img_path = "utils/oefa-logo-header.png"

section = doc.sections[0]
section.left_margin = Cm(Margenes_Informe.izquierdo)    # Margen izquierdo
section.right_margin = Cm(Margenes_Informe.derecho)   # Margen derecho
section.top_margin = Cm(Margenes_Informe.superior)     # Margen superior
section.bottom_margin = Cm(Margenes_Informe.inferior)  # Margen inferior

# Crear el encabezado
header = section.header

# Configurar la tabla en el encabezado
tabla = header.add_table(rows=1, cols=3, width=Cm(16.7))  # Ancho total de la tabla
tabla.style = 'Table Grid'
tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

# Primera celda (imagen)
celda_1 = tabla.cell(0, 0)
celda_1.height = Cm(tabla_cabecera.celda_alto)
celda_1.width = Cm(tabla_cabecera.celda1_ancho)
celda_1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
parrafo_imagen = celda_1.paragraphs[0]
parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
parrafo_imagen.add_run().add_picture(img_path, width=Cm(3.8), height=Cm(0.80))

# Segunda celda (texto centrado)
celda_2 = tabla.cell(0, 1)
celda_2.height = Cm(tabla_cabecera.celda_alto)
celda_2.width = Cm(tabla_cabecera.celda2_ancho)
celda_2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
parrafo_2 = celda_2.paragraphs[0]
parrafo_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
parrafo_2.add_run(tabla_cabecera.texto_mapro)

# Tercera celda (texto pequeño, alineado a la izquierda)
celda_3 = tabla.cell(0, 2)
celda_3.height = Cm(tabla_cabecera.celda_alto)
celda_3.width = Cm(tabla_cabecera.celda3_ancho)
parrafo_3 = celda_3.paragraphs[0]
parrafo_3.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = parrafo_3.add_run(tabla_cabecera.texto_version)
run.font.size = Pt(9)


# Configurar el estilo global del documento
estilo = doc.styles['Normal']
fuente = estilo.font
fuente.name = 'Arial'  # Tipo de letra
fuente.size = Pt(10)   # Tamaño de letra

titulo_parrafo = doc.add_paragraph()
titulo = titulo_parrafo.add_run('INFORME DE ACTIVIDADES')
titulo_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
titulo.font.bold = True

enca1 = doc.add_paragraph()
enca1_valor = enca1.add_run('\tA')
enca1_valor.font.bold = True
enca1_valor = enca1.add_run('\t\t\t:\t')
enca1_valor = enca1.add_run(receptor.director_tercero)

enca2 = doc.add_paragraph()
enca2_valor = enca2.add_run('\tDE')
enca2_valor.font.bold = True
enca2_valor = enca2.add_run('\t\t\t:\t')
enca2_valor = enca2.add_run(emisor.nombre)

enca3 = doc.add_paragraph()
enca3_valor = enca3.add_run('\tNÚMERO DE')
enca3_valor.font.bold = True
enca3_valor = enca3.add_run('\t\t:\t')
enca3_valor = enca3.add_run(Contrato_Adenda()[1])
enca3_valor = enca3.add_run('\n\tCONTRATO')
enca3_valor.font.bold = True

aucnoc = list(fechas_entregables.keys())

enca4 = doc.add_paragraph()
enca4_valor = enca4.add_run('\tPERIODO DE')
enca4_valor.font.bold = True
enca4_valor = enca4.add_run('\t\t:\t')
enca4_valor = enca4.add_run(f'Del {datetime.strftime(datetime.strptime(emisor.inicio_contrato,"%d/%m/%y"),"%d de %B")} al {datetime.strftime(datetime.strptime(aucnoc[-1], "%d/%m/%y"),"%d de %B de %Y")}')
enca4_valor = enca4.add_run('\n\tCONTRATACIÓN')
enca4_valor.font.bold = True

enca5 = doc.add_paragraph()
enca5_valor = enca5.add_run('\tPERIODO EN EL')
enca5_valor.font.bold = True
enca5_valor = enca5.add_run('\n\tQUE SE OTROGA')
enca5_valor.font.bold = True
enca5_valor = enca5.add_run('\t:\t')
enca5_valor = enca5.add_run(f'Del {datetime.strftime(datetime.strptime(fecha_posterior,"%d/%m/%y"),"%d de %B")} al {datetime.strftime(datetime.strptime(entregable_actual,"%d/%m/%y"),"%d de %B de %Y")}')
enca5_valor.font.bold = True
enca5_valor = enca5.add_run('\n\tLA CONFORMIDAD')
enca5_valor.font.bold = True

enca6 = doc.add_paragraph()
enca6_valor = enca6.add_run('\tENTREGABLE')
enca6_valor.font.bold = True
enca6_valor = enca6.add_run('\t\t:\t')
enca6_valor = enca6.add_run(f'{emisor.entregable} entregable')

enca7 = doc.add_paragraph()
enca7_valor = enca7.add_run('\tFECHA')
enca7_valor.font.bold = True
enca7_valor = enca7.add_run('\t\t\t:\t')
enca7_valor = enca7.add_run(f'{datetime.now().strftime("%d/%m/%Y")}')


# Agregar un párrafo con un borde inferior que actúe como línea
p = enca7._element
p_pr = p.get_or_add_pPr()  # Obtener o crear propiedades del párrafo
p_borders = parse_xml(
    r"""
    <w:pBdr %s>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
    </w:pBdr>
    """ % nsdecls('w')
)
p_pr.append(p_borders)

linea_largo = 1.25

# Ajustar las sangrías del párrafo para reducir el largo de la línea
enca7.paragraph_format.left_indent = Cm(linea_largo)  # Sangría izquierda de 3 cm
enca7.paragraph_format.right_indent = Cm(linea_largo)  # Sangría derecha de 3 cm

#Parrafo 1
i = 0
enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}   {orden_llave[i]}')
enca8_valor.font.bold = True

enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'{Orden[orden_llave[i]][0]}')
enca8.paragraph_format.left_indent = Cm(1.75) 
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

enca9 = doc.add_paragraph()

actividades = list(Actividades.values())

for index, actividad in enumerate(actividades):

    
    if index < (len(actividades) - 1):
        enca9_valor = enca9.add_run(f'{actividad}\n\r')
    else:
        enca9_valor = enca9.add_run(f'{actividad}')
    enca9.paragraph_format.left_indent = Cm(2)

enca8 = doc.add_paragraph()  
enca8_valor = enca8.add_run(f'{Orden[orden_llave[i]][1]}')
enca8.paragraph_format.left_indent = Cm(1.75)
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#Parrafo 2

i = 1
sub_i = 1
enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}   {orden_llave[i]}')
enca8_valor.font.bold = True

enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}{sub_i}  {Orden[orden_llave[i]][0]}')
enca8_valor.font.bold = True

Prim_Act = doc.add_paragraph()
Prim_Act.paragraph_format.left_indent = Cm(2)
Prim_Act_valor = Prim_Act.add_run(f'{Actividades["Primera_Actividad"]}')
Prim_Act_valor.font.bold = True
Prim_Act.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

Seg_Act = doc.add_paragraph()
Seg_Act.paragraph_format.left_indent = Cm(2)
Seg_Act_valor = Seg_Act.add_run(f'{Actividades["Segunda_Actividad"]}')
Seg_Act_valor.font.bold = True
Seg_Act.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

Ter_Act = doc.add_paragraph()
Ter_Act.paragraph_format.left_indent = Cm(2)
Ter_Act_valor = Ter_Act.add_run(f'{Actividades["Tercera_Actividad"]}')
Ter_Act_valor.font.bold = True
Ter_Act.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

Cua_Act = doc.add_paragraph()
Cua_Act.paragraph_format.left_indent = Cm(2)
Cua_Act_valor = Cua_Act.add_run(f'{Actividades["Cuarta_Actividad"]}')
Cua_Act_valor.font.bold = True
Cua_Act.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#Parrafo 3
i = 1
sub_i = 2
enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}{sub_i}  {Orden[orden_llave[i]][0]}')
enca8_valor.font.bold = True
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

enca8 = doc.add_paragraph()
enca8.paragraph_format.left_indent = Cm(2)
enca8_valor = enca8.add_run(f'{Rellenos[sub_i-2]}')
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


#Parrafo 4

i = 1
sub_i = 3
enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}{sub_i}  {Orden[orden_llave[i]][0]}')
enca8_valor.font.bold = True
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

enca8 = doc.add_paragraph()
enca8.paragraph_format.left_indent = Cm(2)
enca8_valor = enca8.add_run(f'{Rellenos[sub_i-2]}')
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


#Parrafo 5

i = 1
sub_i = 4
enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}{sub_i}  {Orden[orden_llave[i]][0]}')
enca8_valor.font.bold = True
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

enca8 = doc.add_paragraph()
enca8.paragraph_format.left_indent = Cm(2)
enca8_valor = enca8.add_run(f'{Rellenos[sub_i-2]}')
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



enca8 = doc.add_paragraph()
enca8.paragraph_format.left_indent = Cm(2)
enca8_valor = enca8.add_run(f'{Rellenos[sub_i-1]}')
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


#Parrafo 5
i = 2
enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}   {orden_llave[i]}')
enca8_valor.font.bold = True
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'{Orden[orden_llave[i]][0]}')
enca8.paragraph_format.left_indent = Cm(2) 
enca8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


#Parrafo 6
i = 3
enca8 = doc.add_paragraph()
enca8_valor = enca8.add_run(f'\t{Romano[i]}   {orden_llave[i]}')
enca8_valor.font.bold = True


enca8 = doc.add_paragraph()
enca8.paragraph_format.left_indent = Cm(2) 
enca8_valor = enca8.add_run(f'{Orden[orden_llave[i]][0]}')

parrafo = doc.add_paragraph()
parrafo.paragraph_format.left_indent = Cm(2) 
cuerpo5_text = parrafo.add_run("\n\r\n\r\n\r\n\r\n\r\n\r")
cuerpo6_text = parrafo.add_run("………………………………………………")
cuerpo7_text = parrafo.add_run(f"\n{emisor.nombre}")
cuerpo7_text = parrafo.add_run(f"\n{emisor.cargo}")
cuerpo8_text = parrafo.add_run(f"\nDNI: {emisor.dni}")

# Actividades-
n_sustentos  = iter(Sustentos.keys())

presentacion_sustento = next(n_sustentos)
Prim_Act_valor = Prim_Act.add_run(f'\n\r{presentacion_sustento}.  {Sustentos[presentacion_sustento]["Producto"]}\n\r')
Prim_Act_valor_evi = Prim_Act.add_run(f'Evidencia N° {presentacion_sustento}:\n\r')
Prim_Act_valor_evi.font.bold = True
Prim_Act_valor_evi_1 = Prim_Act.add_run(f'{Sustentos[presentacion_sustento]["Evidencia"]}\n\r')
Prim_Act.paragraph_format.left_indent = Cm(2.2) 

# Guardar el documento
doc.save("documento_con_tabla.docx")